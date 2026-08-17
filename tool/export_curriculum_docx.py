#!/usr/bin/env python3
"""Export the app curriculum to an editable Microsoft Word review document."""

from __future__ import annotations

import argparse
import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = REPO_ROOT / "content" / "drafts" / "lesson_01.json"
DEFAULT_OUTPUT = (
    REPO_ROOT / "docs" / "review" / "arabic-grammar-curriculum-review.docx"
)


def localized(value: dict[str, str]) -> str:
    return f"English: {value['en']}\nArabic: {value['ar']}"


def set_rtl(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    properties = paragraph._p.get_or_add_pPr()
    bidi = properties.find(qn("w:bidi"))
    if bidi is None:
        properties.append(OxmlElement("w:bidi"))


def add_arabic_paragraph(
    container: Any,
    text: str,
    *,
    style: str | None = None,
    bold: bool = False,
) -> Any:
    paragraph = container.add_paragraph(style=style)
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    return paragraph


def shade_cell(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    rtl: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if rtl:
        set_rtl(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bilingual_block(
    document: Document,
    label: str,
    value: dict[str, str],
    *,
    heading_level: int | None = None,
) -> None:
    if heading_level is not None:
        document.add_heading(label, level=heading_level)
    elif label:
        document.add_paragraph(label, style="Heading 4")
    document.add_paragraph(value["en"])
    add_arabic_paragraph(document, value["ar"])


def add_bilingual_table(
    document: Document,
    rows: Iterable[tuple[str, str]],
    *,
    headers: tuple[str, str] = ("English", "Arabic"),
) -> None:
    values = list(rows)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], headers[0], bold=True)
    set_cell_text(table.rows[0].cells[1], headers[1], bold=True, rtl=True)
    shade_cell(table.rows[0].cells[0], "D9EAF7")
    shade_cell(table.rows[0].cells[1], "D9EAF7")
    for english, arabic in values:
        cells = table.add_row().cells
        set_cell_text(cells[0], english)
        set_cell_text(cells[1], arabic, rtl=True)


def add_example(document: Document, example: dict[str, Any]) -> None:
    document.add_heading(
        f"Example: {example['id']}",
        level=4,
    )
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "With vowel marks", bold=True)
    set_cell_text(
        table.cell(0, 1),
        example["vocalized"],
        rtl=True,
    )
    set_cell_text(table.cell(1, 0), "Without vowel marks", bold=True)
    set_cell_text(
        table.cell(1, 1),
        example["unvocalized"],
        rtl=True,
    )

    document.add_paragraph("Token-level analysis", style="Heading 5")
    token_table = document.add_table(rows=1, cols=7)
    token_table.style = "Table Grid"
    headers = ["Token", "Span", "Role", "State", "Sign", "Ending", "Reason"]
    for index, header in enumerate(headers):
        set_cell_text(token_table.rows[0].cells[index], header, bold=True)
        shade_cell(token_table.rows[0].cells[index], "FFF2CC")
    for token in example["tokens"]:
        cells = token_table.add_row().cells
        values = [
            token["text"],
            f"{token['start']}–{token['end']}",
            localized(token["role"]),
            token["grammarState"],
            localized(token["grammaticalSign"]),
            token["ending"],
            localized(token["reason"]),
        ]
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, rtl=index in {0, 5})


def add_exercise(
    document: Document,
    exercise: dict[str, Any],
    number: int,
) -> None:
    document.add_heading(
        f"{number}. {exercise['id']} — {exercise['type']}",
        level=4,
    )
    add_bilingual_block(document, "Prompt", exercise["prompt"])
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Correct", "Option ID", "English label", "Arabic label", "Feedback"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        shade_cell(table.rows[0].cells[index], "E2F0D9")
    for option in exercise["options"]:
        cells = table.add_row().cells
        values = [
            "YES" if option["isCorrect"] else "",
            option["id"],
            option["label"]["en"],
            option["label"]["ar"],
            localized(option["feedback"]),
        ]
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, rtl=index == 3)
        if option["isCorrect"]:
            for cell in cells:
                shade_cell(cell, "E2F0D9")


def add_sources(document: Document, sources: list[dict[str, Any]]) -> None:
    document.add_heading("Sources and provenance", level=3)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Title / Author", "License", "URL", "Citation"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        shade_cell(table.rows[0].cells[index], "D9EAF7")
    for source in sources:
        cells = table.add_row().cells
        values = [
            source["id"],
            f"{localized(source['title'])}\n{localized(source['author'])}",
            source["licenseStatus"],
            source["url"] or "N/A",
            localized(source["citation"]),
        ]
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)


def add_review_worksheet(document: Document, lesson: dict[str, Any]) -> None:
    document.add_heading("Teacher review worksheet", level=3)
    review = lesson["review"]
    document.add_paragraph(
        f"Current app status: {review['status']} | "
        f"Content version: {review['contentVersion']}"
    )
    table = document.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    fields = [
        ("Decision", "Approve / Approve with changes / Return for revision"),
        ("Reviewer name and qualifications", ""),
        ("Review date", ""),
        ("Required corrections", ""),
        ("Optional recommendations", ""),
        ("Final approval signature or initials", ""),
    ]
    for row, (label, value) in zip(table.rows, fields):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
        row.height = Inches(0.45 if label != "Required corrections" else 1.0)


def add_lesson(
    document: Document,
    lesson: dict[str, Any],
    level: dict[str, Any],
) -> None:
    document.add_heading(
        f"Lesson {lesson['order']}: {lesson['title']['en']}",
        level=1,
    )
    add_arabic_paragraph(
        document,
        lesson["title"]["ar"],
        style="Title",
        bold=True,
    )
    document.add_paragraph(
        f"Lesson ID: {lesson['id']} | Level: {level['id']} | "
        f"Estimated time: {lesson['estimatedMinutes']} minutes | "
        f"Prerequisites: {', '.join(lesson['prerequisites']) or 'None'}"
    )

    document.add_heading("Objectives", level=3)
    add_bilingual_table(
        document,
        ((item["en"], item["ar"]) for item in lesson["objectives"]),
    )

    document.add_heading("Teaching sections and examples", level=2)
    for section in lesson["sections"]:
        document.add_heading(
            f"{section['title']['en']} ({section['id']} — {section['type']})",
            level=3,
        )
        document.add_paragraph(section["body"]["en"])
        add_arabic_paragraph(document, section["title"]["ar"], bold=True)
        add_arabic_paragraph(document, section["body"]["ar"])
        for example in section["examples"]:
            add_example(document, example)

    document.add_heading("Primary exercises", level=2)
    for number, exercise in enumerate(lesson["exercises"], start=1):
        add_exercise(document, exercise, number)

    document.add_heading("Alternate repeat exercises", level=2)
    for number, exercise in enumerate(lesson["repeatExercises"], start=1):
        add_exercise(document, exercise, number)

    add_sources(document, lesson["sources"])
    add_review_worksheet(document, lesson)


def configure_document(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(26)
    styles["Title"].font.color.rgb = RGBColor(20, 75, 110)
    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(20, 75, 110)

    for section in document.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)


def add_front_matter(
    document: Document,
    catalog: dict[str, Any],
    source_hash: str,
    generated_at: datetime,
) -> None:
    document.add_heading("I'rab Curriculum Review Document", level=0)
    add_arabic_paragraph(document, "وثيقة مراجعة منهج إعراب", style="Title")
    document.add_paragraph(
        "Editable teacher-review export generated directly from the same "
        "curriculum JSON bundled in the app."
    )
    metadata = [
        ("Curriculum version", catalog["contentVersion"]),
        ("Schema version", str(catalog["schemaVersion"])),
        ("Levels", str(len(catalog["levels"]))),
        ("Lessons", str(len(catalog["lessons"]))),
        ("Generated", generated_at.isoformat()),
        ("Source SHA-256", source_hash),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in metadata:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)

    document.add_heading("Reviewer instructions", level=1)
    for instruction in [
        "Use Word comments for observations tied to exact wording.",
        "Use Track Changes for proposed replacement text.",
        "Review both primary and alternate exercises; both appear in the app.",
        "Verify each marked correct answer and all feedback, including incorrect-answer feedback.",
        "Record required corrections and approval in each lesson's worksheet.",
        "The app content must not be promoted to production until qualified review is complete.",
    ]:
        document.add_paragraph(instruction, style="List Bullet")

    document.add_heading("Curriculum index", level=1)
    lessons_by_id = {lesson["id"]: lesson for lesson in catalog["lessons"]}
    for level in catalog["levels"]:
        document.add_heading(
            f"{level['title']['en']} / {level['title']['ar']}",
            level=2,
        )
        document.add_paragraph(localized(level["description"]))
        for lesson_id in level["lessonIds"]:
            lesson = lessons_by_id[lesson_id]
            document.add_paragraph(
                f"{lesson['order']}. {lesson['title']['en']} / "
                f"{lesson['title']['ar']} ({lesson_id})",
                style="List Number",
            )
    document.add_page_break()


def add_headers_and_footers(document: Document, content_version: str) -> None:
    # Sections inherit a linked header/footer by default. Set the shared content
    # once; appending in a loop would repeat the text for every linked section.
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"I'rab curriculum review — {content_version}"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.text = "Ebaid LLC — Teacher review draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)


def export(source_path: Path, output_path: Path) -> None:
    source_bytes = source_path.read_bytes()
    catalog = json.loads(source_bytes)
    source_hash = hashlib.sha256(source_bytes).hexdigest()
    generated_at = datetime.now(timezone.utc)
    levels_by_lesson = {
        lesson_id: level
        for level in catalog["levels"]
        for lesson_id in level["lessonIds"]
    }

    document = Document()
    configure_document(document)
    add_front_matter(document, catalog, source_hash, generated_at)
    for index, lesson in enumerate(catalog["lessons"]):
        add_lesson(document, lesson, levels_by_lesson[lesson["id"]])
        if index < len(catalog["lessons"]) - 1:
            document.add_section(WD_SECTION.NEW_PAGE)
    add_headers_and_footers(document, catalog["contentVersion"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(
        f"Exported {len(catalog['lessons'])} lessons from "
        f"{catalog['contentVersion']} to {output_path}"
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    export(args.source.resolve(), args.output.resolve())
